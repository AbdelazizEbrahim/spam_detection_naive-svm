from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

def set_font(paragraph, font_name="Calibri", font_size=11):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    set_font(heading, font_size=14 if level == 1 else 12)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    set_font(p)
    return p

def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        set_font(p)

def add_code_block(doc, code):
    p = doc.add_paragraph(code, style='No Spacing')
    set_font(p, font_name="Courier New")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Set headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_font(hdr_cells[0].paragraphs[0], font_size=11, font_name="Calibri")
    # Add rows
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)
            set_font(row_cells[j].paragraphs[0])

def create_documentation():
    doc = Document()

    # Title
    add_heading(doc, "Spam Detection Pipeline Documentation", level=1)

    # Overview
    add_heading(doc, "Overview", level=2)
    add_paragraph(doc, "This Python script implements a spam detection pipeline using two machine learning algorithms: Naïve Bayes and Support Vector Machine (SVM). The pipeline processes a text dataset, trains both models, evaluates their performance, and allows users to interactively classify new text inputs as spam or not_spam. The script includes logging for debugging and performance tracking, and it uses TF-IDF vectorization for text preprocessing.")

    add_heading(doc, "Key Features", level=3)
    add_bullet_list(doc, [
        "Data Preprocessing: Cleans and loads a tab-separated dataset (dataset.csv) with label (spam/ham) and text columns.",
        "Text Vectorization: Uses TF-IDF to convert text into numerical features, with stop-word removal and n-gram support.",
        "Model Training: Trains Naïve Bayes (MultinomialNB) and SVM (LinearSVC) models.",
        "Evaluation: Computes accuracy, precision, recall, and F1-score for both models, with a comparison of results.",
        "Interactive Mode: Allows continuous user input to classify text as spam or not_spam, with feature explanations for predictions.",
        "Logging: Records pipeline progress, timings, and errors to spam_detection.log and the console."
    ])

    # Setup Instructions
    add_heading(doc, "Setup Instructions", level=2)
    add_bullet_list(doc, [
        "Prerequisites:",
        "  - Python 3.8+",
        "  - Required libraries: pandas, numpy, scikit-learn",
        "  - Install dependencies using:"
    ])
    add_code_block(doc, "pip install pandas numpy scikit-learn")
    add_bullet_list(doc, [
        "Dataset Preparation:",
        "  - The pipeline expects a tab-separated dataset.csv file in the data/ directory with two columns: label (spam or ham) and text.",
        "  - Example dataset format:"
    ])
    add_code_block(doc, "ham\tGo until jurong point, crazy.. Available only in bugis n great world la e buffet... Cine there got amore wat...\nspam\tFree entry in 2 a wkly comp to win FA Cup final tkts 21st May 2005. Text FA to 87121 to receive entry question(std txt rate)T&C's apply 08452810075over18's")
    add_bullet_list(doc, [
        "  - Place the dataset at data/dataset.csv. If the file is missing, the script will raise an error.",
        "Running the Script:",
        "  - Save the scripts as naive_classifier.py (Naïve Bayes) and svm_model.py (SVM), or use the combined pipeline.",
        "  - Run either script:"
    ])
    add_code_block(doc, "python naive_classifier.py\npython svm_model.py")
    add_bullet_list(doc, [
        "  - The script will:",
        "    - Clean and load the dataset (5,574 samples: 4,827 not_spam, 747 spam).",
        "    - Train and evaluate the model.",
        "    - Display performance metrics and top features.",
        "    - Enter interactive mode for user input."
    ])

    # Pipeline Components
    add_heading(doc, "Pipeline Components", level=2)
    
    add_heading(doc, "1. Data Preprocessing", level=3)
    add_bullet_list(doc, [
        "Cleaning: Removes empty lines, extra tabs, and quotes from the dataset. Saves cleaned data to data/cleaned_dataset_nb.csv (Naïve Bayes) or data/cleaned_dataset_svm.csv (SVM).",
        "Loading: Reads the dataset into a Pandas DataFrame, maps ham to not_spam, and drops invalid rows.",
        "Output: Logs 5,574 valid samples with class distribution {'not_spam': 4827, 'spam': 747}."
    ])

    add_heading(doc, "2. Text Vectorization", level=3)
    add_bullet_list(doc, [
        "Uses TfidfVectorizer with:",
        "  - English stop-word removal.",
        "  - Maximum 20,000 features.",
        "  - N-grams (1 to 3 words).",
        "  - Minimum document frequency of 2.",
        "Converts training and test text into sparse numerical matrices.",
        "Logs vocabulary size of 9,977 terms."
    ])

    add_heading(doc, "3. Model Training", level=3)
    add_bullet_list(doc, [
        "Naïve Bayes: Uses MultinomialNB with smoothing parameter alpha=0.5. Training takes ~0.0045 seconds.",
        "SVM: Uses LinearSVC with C=1.0, balanced class weights, and maximum 1,000 iterations. Training takes ~0.0138 seconds.",
        "Logs training completion and time."
    ])

    add_heading(doc, "4. Model Evaluation", level=3)
    add_bullet_list(doc, [
        "Computes:",
        "  - Accuracy.",
        "  - Precision, recall, and F1-score for the spam class.",
        "Results from the dataset:"
    ])
    add_paragraph(doc, "Naïve Bayes Results:")
    add_code_block(doc, "--- Naïve Bayes Results ---\nAccuracy: 0.9776\nPrecision (spam): 0.9921\nRecall (spam): 0.8389\nF1-score (spam): 0.9091")
    add_paragraph(doc, "SVM Results:")
    add_code_block(doc, "--- SVM Results ---\nAccuracy: 0.9848\nPrecision (spam): 0.9521\nRecall (spam): 0.9329\nF1-score (spam): 0.9424")

    add_heading(doc, "5. Feature Analysis", level=3)
    add_bullet_list(doc, [
        "Logs the top 10 features contributing to spam detection for both models.",
        "Naïve Bayes: Features with highest spam probabilities:"
    ])
    add_code_block(doc, "Top spam features: [('free', 0.0027), ('txt', 0.0019), ('stop', 0.0017), ('text', 0.0016), ('mobile', 0.0016), ('claim', 0.0015), ('reply', 0.0015), ('ur', 0.0014), ('www', 0.0014), ('prize', 0.0013)]")
    add_bullet_list(doc, [
        "SVM: Features with highest absolute weights:"
    ])
    add_code_block(doc, "Top SVM features: [('uk', 2.2733), ('mobile', 2.1749), ('txt', 1.9388), ('claim', 1.8962), ('150p', 1.7803), ('won', 1.7490), ('50', 1.7381), ('www', 1.6830), ('com', 1.6697), ('video', 1.5602)]")

    add_heading(doc, "6. Interactive Mode", level=3)
    add_bullet_list(doc, [
        "Allows continuous user input to classify text.",
        "For each input:",
        "  - Predicts spam or not_spam using the model.",
        "  - Displays prediction time (~0.001-0.002 seconds).",
        "  - Shows top 5 features influencing the prediction.",
        "Exits when the user types 'exit' or interrupts with Ctrl+C.",
        "Handles empty inputs and errors gracefully."
    ])

    # Demonstration with Example Data
    add_heading(doc, "Demonstration with Example Data", level=2)
    add_paragraph(doc, "Below are interactions from the SMS Spam Collection dataset, showing predictions for five example texts using both Naïve Bayes and SVM models.")

    add_heading(doc, "Example 1: Potential Spam Text", level=3)
    add_paragraph(doc, "Input:")
    add_code_block(doc, "FreeMsg Hey there darling it's been 3 week's now and no word back! I'd like some fun you up for it still? Tb ok! XxX std chgs to send, £1.50 to rcv")
    add_paragraph(doc, "Naïve Bayes Output:")
    add_code_block(doc, "Prediction: not_spam (Time: 0.0018s)\nTop features: [('50', 0.0009), ('send', 0.0008), ('week', 0.0008), ('word', 0.0004), ('freemsg', 0.0003)]")
    add_paragraph(doc, "SVM Output:")
    add_code_block(doc, "Prediction: spam (Time: 0.0019s)\nTop features: [('50', 1.7381), ('std', 0.9097), ('freemsg', 0.8272), ('hey', -0.7746), ('ok', -0.7111)]")
    add_paragraph(doc, "Interpretation:")
    add_bullet_list(doc, [
        "Naïve Bayes incorrectly predicts not_spam, possibly due to conversational tone ('hey', 'ok') outweighing spam indicators.",
        "SVM correctly predicts spam, with strong positive weights for '50', 'std', and 'freemsg' indicating promotional content."
    ])

    add_heading(doc, "Example 2: Not Spam Text", level=3)
    add_paragraph(doc, "Input:")
    add_code_block(doc, "Nah I don't think he goes to usf, he lives around here though")
    add_paragraph(doc, "Naïve Bayes Output:")
    add_code_block(doc, "Prediction: not_spam (Time: 0.0024s)\nTop features: [('don', 0.0002), ('think', 0.0002), ('don think', 0.0001), ('goes', 0.0001), ('lives', 0.0001)]")
    add_paragraph(doc, "SVM Output:")
    add_code_block(doc, "Prediction: not_spam (Time: 0.0016s)\nTop features: [('don', -0.4966), ('nah', -0.2536), ('think', -0.2475), ('goes', -0.1640), ('usf', -0.0944)]")
    add_paragraph(doc, "Interpretation:")
    add_bullet_list(doc, [
        "Both models correctly predict not_spam.",
        "Features like 'don', 'think', and 'nah' have low spam probabilities (Naïve Bayes) or negative weights (SVM), indicating casual conversation."
    ])

    add_heading(doc, "Example 3: Spam Text", level=3)
    add_paragraph(doc, "Input:")
    add_code_block(doc, "Had your mobile 11 months or more? U R entitled to Update to the latest colour mobiles with camera for Free! Call The Mobile Update Co FREE on 08002986030")
    add_paragraph(doc, "Naïve Bayes Output:")
    add_code_block(doc, "Prediction: spam (Time: 0.0020s)\nTop features: [('free', 0.0027), ('mobile', 0.0016), ('latest', 0.0005), ('camera', 0.0005), ('colour', 0.0003)]")
    add_paragraph(doc, "SVM Output:")
    add_code_block(doc, "Prediction: spam (Time: 0.0020s)\nTop features: [('mobile', 2.1749), ('latest', 0.8945), ('free', 0.8280), ('camera', 0.7550), ('mobiles', 0.3984)]")
    add_paragraph(doc, "Interpretation:")
    add_bullet_list(doc, [
        "Both models correctly predict spam.",
        "Key features like 'free', 'mobile', and 'latest' have high spam probabilities (Naïve Bayes) or positive weights (SVM), indicating promotional content."
    ])

    add_heading(doc, "Example 4: Not Spam Text", level=3)
    add_paragraph(doc, "Input:")
    add_code_block(doc, "Is that seriously how you spell his name?")
    add_paragraph(doc, "Naïve Bayes Output:")
    add_code_block(doc, "Prediction: not_spam (Time: 0.0013s)\nTop features: [('seriously', 0.0001), ('spell', 0.0001)]")
    add_paragraph(doc, "SVM Output:")
    add_code_block(doc, "Prediction: not_spam (Time: 0.0019s)\nTop features: [('seriously', -0.1100), ('spell', -0.1066)]")
    add_paragraph(doc, "Interpretation:")
    add_bullet_list(doc, [
        "Both models correctly predict not_spam.",
        "Features 'seriously' and 'spell' have low spam probabilities (Naïve Bayes) or negative weights (SVM), indicating non-promotional content."
    ])

    add_heading(doc, "Example 5: Spam Text", level=3)
    add_paragraph(doc, "Input:")
    add_code_block(doc, "Thanks for your subscription to Ringtone UK your mobile will be charged £5/month Please confirm by replying YES or NO. If you reply NO you will not be charged")
    add_paragraph(doc, "Naïve Bayes Output:")
    add_code_block(doc, "Prediction: spam (Time: 0.0019s)\nTop features: [('mobile', 0.0016), ('reply', 0.0015), ('uk', 0.0012), ('ringtone', 0.0007), ('charged', 0.0004)]")
    add_paragraph(doc, "SVM Output:")
    add_code_block(doc, "Prediction: spam (Time: 0.0009s)\nTop features: [('uk', 2.2733), ('mobile', 2.1749), ('reply', 1.5506), ('ringtone', 1.4442), ('charged', 0.9049)]")
    add_paragraph(doc, "Interpretation:")
    add_bullet_list(doc, [
        "Both models correctly predict spam.",
        "Features like 'uk', 'mobile', and 'reply' have high spam probabilities (Naïve Bayes) or positive weights (SVM), indicating a subscription-based scam."
    ])

    # Usage Notes
    add_heading(doc, "Usage Notes", level=2)
    add_bullet_list(doc, [
        "Dataset: Ensure data/dataset.csv exists and follows the expected format. The dataset used has 5,574 samples (4,827 not_spam, 747 spam).",
        "Performance: SVM (98.48% accuracy) outperforms Naïve Bayes (97.76% accuracy), with higher recall (0.9329 vs. 0.8389) and F1-score (0.9424 vs. 0.9091) for spam.",
        "Feature Explanations: Top features provide insight into predictions. For Naïve Bayes, higher probabilities indicate spam. For SVM, positive weights indicate spam, negative weights indicate not_spam.",
        "Error Handling: The script logs errors to naive_bayes.log or svm.log and handles invalid inputs in interactive mode."
    ])

    # Limitations
    add_heading(doc, "Limitations", level=2)
    add_bullet_list(doc, [
        "The dataset is imbalanced (13.4% spam), which may affect Naïve Bayes recall (0.8389). Oversampling or class weighting could improve performance.",
        "The TF-IDF vectorizer is limited to 20,000 features, resulting in 9,977 terms. Rare terms may be excluded.",
        "Interactive mode is terminal-based. A GUI would require additional libraries (e.g., tkinter)."
    ])

    # Future Improvements
    add_heading(doc, "Future Improvements", level=2)
    add_bullet_list(doc, [
        "Implement cross-validation to improve model robustness.",
        "Add support for other algorithms (e.g., Random Forest, Neural Networks).",
        "Develop a web interface for user input.",
        "Include hyperparameter tuning for Naïve Bayes and SVM."
    ])

    # License
    add_heading(doc, "License", level=2)
    add_paragraph(doc, "This project is for educational purposes and provided as-is. The SMS Spam Collection dataset is publicly available for research.")

    # Footer
    add_paragraph(doc, "Generated on May 06, 2025").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    doc.save('spam_detection_documentation.docx')
    print("Documentation saved as spam_detection_documentation.docx")

if __name__ == "__main__":
    create_documentation()